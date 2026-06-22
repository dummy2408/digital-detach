import os
from docx import Document
from docx.shared import Inches

def append_images(folder_name):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    report_path = os.path.join(base_dir, folder_name, f'Evaluation_Report_{folder_name.replace("Regressor", "")}.docx')
    eval_dir = os.path.join(base_dir, folder_name, 'evaluation')
    
    if not os.path.exists(report_path):
        print(f"Report not found at {report_path}")
        return
        
    doc = Document(report_path)
    
    # Model Performance
    doc.add_heading('Model Performance', level=1)
    perf_images = [
        ('Confusion Matrix', 'model_performance/confusion_matrix.png'),
        ('ROC Curves', 'model_performance/roc_curves.png'),
        ('Precision-Recall Curves', 'model_performance/precision_recall.png'),
        ('Calibration Curve', 'model_performance/calibration_curve.png')
    ]
    
    for title, img_path in perf_images:
        full_path = os.path.join(eval_dir, img_path)
        if os.path.exists(full_path):
            doc.add_heading(title, level=2)
            doc.add_picture(full_path, width=Inches(5.5))
            
    # Model Analysis
    doc.add_heading('Model Analysis', level=1)
    analysis_images = [
        ('Feature Importance', 'model_analysis/feature_importance.png'),
        ('Learning Curves', 'model_analysis/learning_curves.png'),
        ('Correlation Heatmap', 'model_analysis/correlation_heatmap.png')
    ]
    
    for title, img_path in analysis_images:
        full_path = os.path.join(eval_dir, img_path)
        if os.path.exists(full_path):
            doc.add_heading(title, level=2)
            doc.add_picture(full_path, width=Inches(5.5))
            
    doc.save(report_path)
    print(f"Updated {report_path} with images.")

if __name__ == '__main__':
    append_images("RandomForestRegressor")
    append_images("DecisionTreeRegressor")
