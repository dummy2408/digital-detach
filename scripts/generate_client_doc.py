import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('How to Run Digital Detach', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Follow these exact steps to start the AI engine and connect it to the live website. The frontend is already hosted online, so you only need to turn on the backend connection.')

# Step 1
doc.add_heading('Step 1: Prepare the Laptop', level=2)
p1 = doc.add_paragraph(style='List Number')
p1.add_run('Download and install Python from python.org.\n')
p1.add_run('CRITICAL: During installation, you MUST check the box at the very bottom that says "Add python.exe to PATH" before clicking Install.').bold = True

p2 = doc.add_paragraph('Extract the project Zip file onto your Desktop.', style='List Number')

# Step 2
doc.add_heading('Step 2: Turn on the Backend Engine', level=2)
p3 = doc.add_paragraph('Open the extracted folder, and go inside the "backend" folder.', style='List Number')
p4 = doc.add_paragraph('Click on the folder\'s address bar at the very top (where it says C:\\Users\\...\\backend), delete the text, type exactly "cmd", and hit Enter. A black window will pop up.', style='List Number')
p5 = doc.add_paragraph('In that black window, type this exact command and hit enter to install the AI tools (this takes a few minutes):\n', style='List Number')
p5.add_run('pip install -r requirements.txt').bold = True

p6 = doc.add_paragraph('Once it finishes, type this command and hit enter to turn on the engine:\n', style='List Number')
p6.add_run('uvicorn api.index:app --port 8000').bold = True
p7 = doc.add_paragraph('(Leave this black window open!)')
p7.runs[0].italic = True

# Step 3
doc.add_heading('Step 3: Turn on the Internet Bridge (Ngrok)', level=2)
p8 = doc.add_paragraph('Open a new "cmd" black window by clicking the address bar of the folder again and typing "cmd".', style='List Number')
p9 = doc.add_paragraph('First, install Ngrok by typing this command and hitting Enter (if it asks for permission, type Y and hit Enter):\n', style='List Number')
p9.add_run('winget install ngrok').bold = True

p10 = doc.add_paragraph('IMPORTANT: Once it is installed, CLOSE that black window and open a BRAND NEW "cmd" black window. (This is required so the laptop recognizes the new ngrok command).', style='List Number')

p11 = doc.add_paragraph('Now, log in by pasting this command and hitting Enter:\n', style='List Number')
p11.add_run('ngrok config add-authtoken 3FJ6zKtq9Zi54gPVuV6Upv3aDZH_79fhMXAg5LCnkGL6vjQKa').bold = True

p12 = doc.add_paragraph('Finally, turn on the bridge by running this command:\n', style='List Number')
p12.add_run('ngrok http --domain=dislike-film-unaudited.ngrok-free.dev 8000').bold = True
p13 = doc.add_paragraph('(Leave this black window open too!)')
p13.runs[0].italic = True

# Step 4
doc.add_heading('Step 4: You are live!', level=2)
doc.add_paragraph('Everything is now running! Anyone in the world can go to your permanent website link (https://digital-detach.vercel.app/) on their phone or computer, and it will magically connect to your laptop to process the AI predictions.')

# Save to docs folder
doc_path = 'd:/Digital-Detach/docs/Client_Setup_Guide.docx'
doc.save(doc_path)
print(f'Document generated successfully at: {doc_path}')
