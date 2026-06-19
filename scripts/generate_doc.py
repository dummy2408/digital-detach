import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Digital Detach: Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Digital Detach is an AI-powered digital wellness application designed to help users monitor, "
        "understand, and reduce their smartphone usage. The goal of the project is to provide actionable "
        "insights by analyzing screen time screenshots using computer vision, predicting a user's addiction "
        "risk level using a trained machine learning model, and assigning them a behavioral persona through "
        "clustering. Finally, it provides rule-based recommendations and uses a gamified 'living plant' "
        "visualization to encourage healthier digital habits."
    )
    
    # 2. Project Explanation (Step-by-Step)
    doc.add_heading('2. Project Explanation (Step-by-Step)', level=1)
    
    steps = [
        ("Step 1: User Signup and Survey", 
         "What happens: The user signs up and completes a 10-question initial profile survey.\n"
         "The 10 questions are:\n"
         "1. Full Name\n"
         "2. Age\n"
         "3. Gender\n"
         "4. School Grade\n"
         "5. Average Sleep (Hours)\n"
         "6. Daily Exercise (Hours)\n"
         "7. Stress & Anxiety Level (1-10)\n"
         "8. Favorite Hobby\n"
         "9. Peak Usage Time (Morning, Afternoon, Night)\n"
         "10. Usual Session Style (Long uninterrupted vs. Frequent short checks)\n"
         "Why it happens: To gather baseline personal data. This baseline is later combined with the screenshot data by the model to provide a more accurate risk prediction every time the user analyzes their screen time.\n"
         "Where it is implemented: Frontend in 'frontend/src/app/components/ProfileSetup.tsx' and saved to the Supabase 'profiles' table."),
        
        ("Step 2: User Image Upload and Data Extraction", 
         "What happens: The user uploads a screenshot of their screen time. The app extracts text and metrics from the image.\n"
         "Why it happens: To automatically gather usage data without requiring manual data entry.\n"
         "Where it is implemented: Frontend in 'frontend/src/app/components/Dashboard.tsx' using the Google Gemini Vision API to parse the image into structured JSON fields."),
        
        ("Step 3: Sending Data to the Backend", 
         "What happens: The extracted metrics (like daily usage, sleep hours, pickup frequency) are sent to the backend server.\n"
         "Why it happens: To process the raw data securely and run complex machine learning predictions.\n"
         "Where it is implemented: Frontend ('Dashboard.tsx') makes an HTTP POST request to the FastAPI backend at 'backend/api/index.py'."),
        
        ("Step 4: Machine Learning Prediction (Risk Level)", 
         "What happens: The backend feeds the data into a trained XGBoost model to predict the user's addiction risk level (e.g., Low, Moderate, High).\n"
         "Why it happens: To classify the user's behavior based on patterns learned from a dataset.\n"
         "Where it is implemented: The model was trained in 'ml/xgboost_model/train_model.py' and is loaded and executed in 'backend/api/index.py'."),
        
        ("Step 5: Handling Data Imbalance (During Training)", 
         "What happens: Synthetic data was generated to balance the risk categories.\n"
         "Why it happens: The original dataset had unequal numbers of low, moderate, and high risk profiles. Without balancing, the model would be biased toward the majority class.\n"
         "Where it is implemented: In 'ml/xgboost_model/train_model.py' using the SMOTE (Synthetic Minority Over-sampling Technique) library."),
        
        ("Step 6: Clustering and Persona Assignment", 
         "What happens: The user is grouped into a specific behavioral cluster and assigned a persona (like 'Heavy Social Media User' or 'Productive/Work Focused').\n"
         "Why it happens: To provide users with a relatable summary of their digital lifestyle rather than just raw numbers.\n"
         "Where it is implemented: A KMeans clustering model was trained alongside XGBoost in 'train_model.py'. It is applied to new users in 'backend/api/index.py'."),
        
        ("Step 7: Explaining the Prediction (SHAP)", 
         "What happens: The system identifies the top 3 specific factors (e.g., 'Too much gaming' or 'High phone pickups') that drove the user's risk score.\n"
         "Why it happens: To build trust and explain the 'black box' machine learning model to the user.\n"
         "Where it is implemented: In 'backend/api/index.py' using the SHAP (SHapley Additive exPlanations) library on the XGBoost TreeExplainer."),
        
        ("Step 8: Rule-Based Recommendations", 
         "What happens: Custom, actionable advice is generated based on the user's specific metrics.\n"
         "Why it happens: Machine learning gives a score, but users need concrete steps to improve. The rule-based logic checks specific thresholds (like 'if sleep < 6 hours' or 'if risk is High').\n"
         "Where it is implemented: In 'backend/api/index.py' within the 'generate_recommendations()' function."),
        
        ("Step 9: Gamification and Database Storage", 
         "What happens: The results are displayed, a visual 'plant' grows or shrinks based on the risk level, and the history is saved.\n"
         "Why it happens: Gamification motivates the user, and history tracking shows progress over time.\n"
         "Where it is implemented: The plant is rendered in 'frontend/src/app/components/PlantVisualization.tsx', and the history is saved to the database in 'Dashboard.tsx'.")
    ]
    
    for i, (title_text, desc_text) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title_text}").bold = True
        doc.add_paragraph(desc_text, style='List Bullet')

    # 3. Tech Stack Section
    doc.add_heading('3. Tech Stack Section', level=1)
    techs = [
        ("Frontend", "Next.js 15, React 19, Tailwind CSS (for styling), Framer Motion (for animations). Next.js provides the fast, modern web framework, while Tailwind allows for rapid and beautiful design."),
        ("Backend", "FastAPI (Python framework), Uvicorn (server). FastAPI is used because it is incredibly fast and highly compatible with Python-based machine learning libraries."),
        ("Machine Learning", "XGBoost (for the core classification model), Scikit-Learn (for KMeans clustering and data scaling), Imbalanced-learn (for SMOTE data balancing), and SHAP (for model explainability)."),
        ("Database", "Supabase (PostgreSQL). Used to store the history of user predictions and whether they found recommendations helpful."),
        ("AI / Computer Vision", "Google Gemini 1.5 Flash. Used specifically for its vision capabilities to extract structured data from user screenshots.")
    ]
    for name, desc in techs:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
        
    # 4. Evaluation Results
    doc.add_heading('4. Evaluation Results', level=1)
    doc.add_paragraph("The following graphs demonstrate the quality of the dataset, the analysis of the XGBoost model, and its overall performance metrics.")
    
    eval_dir = r"d:\Digital-Detach\ml\xgboost_model\evaluation"
    
    sections = {
        "Data Quality": os.path.join(eval_dir, "data_quality"),
        "Model Analysis": os.path.join(eval_dir, "model_analysis"),
        "Model Performance": os.path.join(eval_dir, "model_performance")
    }
    
    for section_name, folder_path in sections.items():
        doc.add_heading(section_name, level=2)
        if os.path.exists(folder_path):
            for file in sorted(os.listdir(folder_path)):
                if file.endswith(".png"):
                    img_path = os.path.join(folder_path, file)
                    try:
                        doc.add_picture(img_path, width=Inches(5.0))
                        caption = doc.add_paragraph(f"Figure: {file.replace('.png', '').replace('_', ' ').title()}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        doc.add_paragraph(f"[Image {file} could not be loaded: {str(e)}]")
        else:
            doc.add_paragraph(f"[Directory not found: {folder_path}]")
            
    # 5. Deployment Strategy
    doc.add_heading('5. Deployment Strategy', level=1)
    doc.add_paragraph("The deployment architecture of Digital Detach is split into two parts: Frontend and Backend.")
    
    doc.add_heading('Frontend Deployment (Vercel)', level=2)
    doc.add_paragraph("The Next.js frontend is deployed globally on Vercel. Vercel automatically builds and serves the Next.js application, providing a fast and secure HTTPS domain (e.g., https://digital-detach.vercel.app/).")

    doc.add_heading('Backend Deployment (Local + Ngrok)', level=2)
    doc.add_paragraph("The FastAPI backend, which requires Python and machine learning libraries like XGBoost, is hosted locally on the developer's machine using Uvicorn on port 8000. To allow the Vercel frontend to communicate with this local backend securely, Ngrok is used as a secure internet tunnel. Ngrok creates a permanent, public HTTPS URL (e.g., https://dislike-film-unaudited.ngrok-free.dev) that forwards all requests to the local port 8000. The frontend's environment variable NEXT_PUBLIC_API_URL is configured to point to this Ngrok URL, enabling seamless integration between the cloud-hosted frontend and locally-hosted AI backend.")
            
    # 6. Q&A Section
    doc.add_heading('6. Q&A Section', level=1)
    doc.add_paragraph("Common questions a reviewer or interviewer might ask regarding this project's architecture and decisions:")
    
    qas = [
        ("Why was XGBoost chosen over other models like Random Forest or standard Neural Networks?",
         "XGBoost is an advanced implementation of gradient boosted decision trees. It generally outperforms Random Forests in tabular data scenarios (like our user metrics) because it builds trees sequentially to correct the errors of previous trees. Neural networks would require significantly more data and computational power, whereas XGBoost is highly efficient and accurate for this dataset scale."),
        
        ("Where in the pipeline does optimization/tuning happen? Is the model optimized?",
         "Model optimization happens in the 'ml/xgboost_model/train_model.py' script. The model was optimized using GridSearchCV (or RandomizedSearchCV) from scikit-learn, which tested various hyperparameter combinations (like learning rate, max depth, and number of estimators) to find the configuration that yielded the best accuracy and F1 score."),
        
        ("Why was KMeans clustering chosen for the persona generation?",
         "KMeans is an unsupervised learning algorithm that excels at partitioning data into distinct, non-overlapping groups based on feature similarity. It was ideal for taking user metrics (screen time, sleep, pickups) and finding natural groupings without needing labeled persona data. This allowed us to organically define personas like 'Heavy Social Media User'."),
        
        ("How was data imbalance addressed and why that method?",
         "The dataset had imbalanced risk levels. This was addressed using SMOTE (Synthetic Minority Over-sampling Technique) in 'train_model.py'. Instead of just duplicating existing data (which leads to overfitting), SMOTE creates synthetic data points by interpolating between existing minority class examples, allowing the model to learn broader patterns for the underrepresented risk levels."),
        
        ("How does the app extract data from the screenshots?",
         "The frontend ('Dashboard.tsx') takes the uploaded image and sends it directly to the Google Gemini 1.5 Flash Vision API using a system prompt that enforces a JSON response. Gemini acts as an advanced OCR, extracting specific metrics (sleep, gaming time) and ignoring irrelevant visual noise."),
        
        ("Why use a separate Python backend instead of Next.js API routes?",
         "Next.js API routes run on Node.js, which is not designed for heavy machine learning tasks. Our backend relies on Python libraries like XGBoost, scikit-learn, and SHAP, which cannot be natively run in Node.js. Therefore, a separate Python FastAPI backend was required to handle the ML pipeline efficiently."),
        
        ("How does the 'rule-based' recommendation system work alongside the ML model?",
         "The ML model only predicts the overall 'Risk Level' (Low, Moderate, High). It does not tell the user *what* to do. The rule-based logic in 'backend/api/index.py' takes the user's specific data inputs (e.g., 'sleep < 6 hours' or 'social media > 3 hours') and triggers specific, pre-written, actionable advice. They complement each other: ML assesses the situation, rules provide the solution."),
        
        ("What is SHAP and why is it used?",
         "SHAP (SHapley Additive exPlanations) is a game-theoretic approach to explain the output of machine learning models. It is used in 'backend/api/index.py' to break down the XGBoost prediction and explicitly show the user which specific features (e.g., 'High pickup frequency') contributed most to their high risk score, ensuring the AI isn't just a 'black box'."),
        
        ("Why did you use FastAPI over Flask or Django?",
         "FastAPI is significantly faster than Flask and provides automatic asynchronous support out of the box. More importantly, it automatically uses Pydantic for data validation, ensuring that the frontend sends correctly typed data to our ML models without writing extensive manual validation code."),
        
        ("How is the database structured?",
         "The app uses a Supabase (PostgreSQL) database. It primarily features a 'history' table that stores user IDs, prediction dates, extracted metrics, the predicted risk level, the assigned cluster/persona, SHAP features, and whether the user found the recommendations helpful (helpful_recs array)."),
        
        ("How does the gamification (Plant Visualization) work?",
         "In 'PlantVisualization.tsx', the component uses CSS and Framer Motion to render a plant that dynamically changes its state. If the user's risk level is 'LOW RISK', the plant renders in full health (blooming, vibrant colors). If the risk is 'HIGH RISK', it visually wilts, using animations and color shifts to provide an emotional visual cue to the user."),
        
        ("Is the user's image saved to your server?",
         "No. The image is passed to the Gemini API as a base64 encoded string directly from the client side or temporarily through the backend, and is immediately discarded after the data is extracted. Only the extracted numerical metrics are stored in the database."),
        
        ("What happens if the Gemini Vision API fails or hallucinates?",
         "The prompt used for Gemini strictly mandates a JSON output with specific keys. If the API fails or returns invalid data, the frontend parsing (`JSON.parse`) catches the error, and the user is prompted to try again or input the data manually."),
        
        ("How do you handle scaling the data for the model?",
         "During training ('train_model.py'), a StandardScaler is fitted to the training data to normalize the features (mean of 0, standard deviation of 1). This scaler is saved as a '.pkl' file and loaded in 'backend/api/index.py' so that any new incoming user data is scaled identically before prediction."),
        
        ("Why are the 'Helpful' and 'Not Helpful' buttons important?",
         "These buttons provide a feedback loop. When a user flags a recommendation as helpful, this preference is stored in the Supabase database. In the future, this data can be used to analyze which rule-based recommendations are effective and potentially train a collaborative filtering recommendation engine."),
        
        ("What was the biggest challenge in building the ML pipeline?",
         "Handling data formatting between the frontend, the FastAPI backend, and the XGBoost model. The model requires an exact array shape with specific feature names in a specific order. Writing the Pydantic model in FastAPI ensured that the data matched exactly what the `.pkl` model expected."),
        
        ("How does the app deal with missing data if the screenshot doesn't have it?",
         "The Gemini prompt instructs the AI to return 0 for any metric it cannot find. The frontend also provides fallback default values when parsing the JSON, ensuring the backend always receives a complete feature vector to pass to the XGBoost model."),
        
        ("Why use Tailwind CSS instead of standard CSS?",
         "Tailwind provides utility classes that allow for rapid, responsive layout development without constantly switching between CSS and TSX files. It was used to quickly implement the dark mode aesthetic, glassmorphism effects, and responsive grid layouts required for the dashboard."),
        
        ("How does the routing work in the Next.js frontend?",
         "The project uses the new Next.js 'App Router' (the `app/` directory). Pages like the authentication screen and dashboard are structured within this routing paradigm, allowing for server-side layout optimizations and simpler nested routing."),
        
        ("What would you add if you had more time?",
         "I would implement OAuth (Google/Apple login) for smoother onboarding, add a feature to continuously poll screen time via an iOS/Android native app rather than requiring screenshots, and train an NLP model to generate dynamic recommendations rather than using rule-based logic.")
    ]
    
    for i, (q, a) in enumerate(qas, 1):
        p_q = doc.add_paragraph()
        p_q.add_run(f"Q{i}: {q}").bold = True
        doc.add_paragraph(a)
        
    os.makedirs(r"d:\Digital-Detach\docs", exist_ok=True)
    doc.save(r"d:\Digital-Detach\docs\Digital_Detach_Documentation.docx")
    print("Document successfully generated at d:\\Digital-Detach\\docs\\Digital_Detach_Documentation.docx")

if __name__ == "__main__":
    create_document()
