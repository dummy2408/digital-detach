from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Helper: shade a paragraph cell-like (for code blocks) ────────────────────
def add_code_block(doc, code_text):
    """Add a shaded paragraph that looks like a code block."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_note(doc, text, color=RGBColor(0x8B, 0x45, 0x13)):
    p = doc.add_paragraph()
    run = p.add_run(f"ℹ  {text}")
    run.font.color.rgb = color
    run.font.italic = True
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_warning(doc, text):
    return add_note(doc, text, color=RGBColor(0xC0, 0x39, 0x2B))

def hr(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── TITLE ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('🌿 Digital Detach — Local Setup Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = sub.add_run('The frontend is already live at ')
r1.font.size = Pt(11)
r2 = sub.add_run('https://digital-detach.vercel.app/')
r2.font.size = Pt(11)
r2.font.bold = True
r2.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
r3 = sub.add_run('\nYou only need to run the backend on your machine.')
r3.font.size = Pt(11)

hr(doc)

# ── HOW IT WORKS ──────────────────────────────────────────────────────────────
doc.add_heading('How This Project Works', level=1)
doc.add_paragraph(
    'The frontend is deployed on Vercel and always online. '
    'The backend (ML model + AI) runs locally on your machine. '
    'Ngrok creates a permanent public HTTPS tunnel so the live website can talk to your local backend.'
)
add_code_block(doc,
    'Your Browser / Phone\n'
    '       ↓\n'
    'https://digital-detach.vercel.app        ← Vercel (always on)\n'
    '       ↓  API calls\n'
    'https://dislike-film-unaudited.ngrok-free.dev  ← Ngrok tunnel\n'
    '       ↓  forwards to\n'
    'http://localhost:8000                    ← FastAPI on YOUR machine'
)

hr(doc)

# ── PREREQUISITES ─────────────────────────────────────────────────────────────
doc.add_heading('Prerequisites', level=1)
doc.add_paragraph('Before starting, make sure you have the following installed:')

pre = [
    ('Python 3.10+',  'Download from https://python.org/downloads\n'
                      'CRITICAL: During installation check "Add python.exe to PATH" before clicking Install.'),
    ('Git',           'Download from https://git-scm.com'),
    ('Ngrok',         'No account needed, no sign up required.\n'
                      'You will install it with a single command in Step 5.'),
]
for name, desc in pre:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name} — ').bold = True
    p.add_run(desc)

hr(doc)

# ── STEP 1 ────────────────────────────────────────────────────────────────────
doc.add_heading('Step 1 — Clone the Repository', level=1)
doc.add_paragraph('Open CMD (or Terminal on Mac) and run:')
add_code_block(doc, 'git clone https://github.com/dummy2408/digital-detach.git\ncd digital-detach')

hr(doc)

# ── STEP 2 ────────────────────────────────────────────────────────────────────
doc.add_heading('Step 2 — Set Up the Credentials', level=1)
add_warning(doc, 'The .env.local file is NOT in the repo for security reasons. You must create it manually.')
doc.add_paragraph(
    'Inside the cloned folder, open the frontend folder. '
    'Create a new file called exactly   .env.local   (with a dot at the start) '
    'and paste the values from the CREDENTIALS.txt file you received into it:'
)
add_code_block(doc,
    'GEMINI_API_KEY=<from CREDENTIALS.txt>\n'
    'NEXT_PUBLIC_SUPABASE_URL=<from CREDENTIALS.txt>\n'
    'NEXT_PUBLIC_SUPABASE_ANON_KEY=<from CREDENTIALS.txt>\n'
    'NEXT_PUBLIC_API_URL=https://dislike-film-unaudited.ngrok-free.dev'
)
add_note(doc, 'CREDENTIALS.txt was shared with you separately (WhatsApp / email). Copy the values exactly as they appear.')

hr(doc)

# ── STEP 3 ────────────────────────────────────────────────────────────────────
doc.add_heading('Step 3 — Install Backend Dependencies', level=1)
doc.add_paragraph('Open CMD inside the backend folder and run:')
add_code_block(doc, 'cd backend\npip install -r requirements.txt')
add_note(doc, 'This installs FastAPI, XGBoost, SHAP, Pandas, Gemini SDK and more. May take a few minutes the first time.')

hr(doc)

# ── STEP 4 ────────────────────────────────────────────────────────────────────
doc.add_heading('Step 4 — Start the Backend Server', level=1)
doc.add_paragraph('From inside the backend folder, run:')
add_code_block(doc, 'uvicorn api.index:app --port 8000')
doc.add_paragraph('You should see:')
add_code_block(doc, 'INFO:     Uvicorn running on http://127.0.0.1:8000 (Press CTRL+C to quit)')
add_warning(doc, 'Keep this CMD window open. The backend stops if you close it.')

hr(doc)

# ── STEP 5 ────────────────────────────────────────────────────────────────────
doc.add_heading('Step 5 — Start the Ngrok Tunnel', level=1)
doc.add_paragraph(
    'Open a BRAND NEW CMD window (keep the backend one running). '
    'Run these 3 commands one by one:'
)
add_note(doc, 'No ngrok account or sign up needed — just copy and paste the commands below.')

doc.add_heading('Command 1 — Install Ngrok', level=2)
doc.add_paragraph('On Windows:')
add_code_block(doc, 'winget install ngrok')
add_warning(doc,
    'After it finishes, CLOSE this CMD window completely and open a fresh new one. '
    'Windows will not recognise the ngrok command until you do this.'
)
doc.add_paragraph('On Mac:')
add_code_block(doc, 'brew install ngrok')

doc.add_heading('Command 2 — Activate the Auth Token', level=2)
doc.add_paragraph('In the new CMD window, run this exactly:')
add_code_block(doc, 'ngrok config add-authtoken 3FJ6zKtq9Zi54gPVuV6Upv3aDZH_79fhMXAg5LCnkGL6vjQKa')
add_note(doc, 'You should see:  Authtoken saved to configuration file')

doc.add_heading('Command 3 — Start the Tunnel', level=2)
add_code_block(doc, 'ngrok http --domain=dislike-film-unaudited.ngrok-free.dev 8000')
doc.add_paragraph('You should see:')
add_code_block(doc, 'Forwarding   https://dislike-film-unaudited.ngrok-free.dev -> http://localhost:8000')
add_warning(doc, 'Keep this CMD window open. The tunnel stops the moment you close it.')

hr(doc)

# ── STEP 6 ────────────────────────────────────────────────────────────────────
doc.add_heading("Step 6 — You're Live! 🎉", level=1)
doc.add_paragraph('Once both the backend and ngrok are running:')
steps = [
    'Go to https://digital-detach.vercel.app/ on any device',
    'The website will connect to your machine\'s backend through the ngrok tunnel',
    'All AI predictions, ML models, and screenshot extraction will work end-to-end',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

hr(doc)

# ── RECAP TABLE ───────────────────────────────────────────────────────────────
doc.add_heading('Recap — What Should Be Running', level=1)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'What'
hdr[1].text = 'Where'
hdr[2].text = 'Command'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows_data = [
    ('FastAPI backend',  'CMD window 1  (inside backend/)', 'uvicorn api.index:app --port 8000'),
    ('Ngrok tunnel',     'CMD window 2  (anywhere)',         'ngrok http --domain=dislike-film-unaudited.ngrok-free.dev 8000'),
    ('Frontend',         'Already live on Vercel',           '— nothing to run —'),
]
for i, (w, wh, cmd) in enumerate(rows_data, start=1):
    row = table.rows[i].cells
    row[0].text = w
    row[1].text = wh
    row[2].text = cmd

hr(doc)

# ── TROUBLESHOOTING ───────────────────────────────────────────────────────────
doc.add_heading('Troubleshooting', level=1)
issues = [
    ('pip not found',                    'Make sure Python was installed with "Add to PATH" checked'),
    ('ngrok: command not found',         'Close and reopen CMD after installing ngrok'),
    ('ModuleNotFoundError on startup',   'Run:  pip install -r requirements.txt  again'),
    ('Website shows backend unavailable','Make sure BOTH CMD windows (uvicorn + ngrok) are still running'),
    ('shap install fails',               'Try:  pip install shap --no-build-isolation'),
]
table2 = doc.add_table(rows=len(issues)+1, cols=2)
table2.style = 'Table Grid'
h = table2.rows[0].cells
h[0].text = 'Problem'
h[1].text = 'Fix'
for cell in h:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for i, (prob, fix) in enumerate(issues, start=1):
    table2.rows[i].cells[0].text = prob
    table2.rows[i].cells[1].text = fix

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = 'd:/Digital-Detach/docs/Setup_Guide.docx'
doc.save(out)
print(f'Done -> {out}')
